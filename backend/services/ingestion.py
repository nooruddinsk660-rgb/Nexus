from __future__ import annotations
from pathlib import Path
from typing import Union
import asyncio, logging, uuid

from models.schemas import Chunk, Modality

log = logging.getLogger("nexus.ingestion")

SUPPORTED: dict[str, Modality] = {
    ".pdf":  Modality.text,
    ".docx": Modality.text,
    ".doc":  Modality.text,
    ".txt":  Modality.text,
    ".png":  Modality.image,
    ".jpg":  Modality.image,
    ".jpeg": Modality.image,
    ".webp": Modality.image,
    ".mp3":  Modality.audio,
    ".mp4":  Modality.audio,
    ".wav":  Modality.audio,
    ".m4a":  Modality.audio,
    ".ogg":  Modality.audio,
}


async def ingest_file(path: Path, file_id: str | None = None) -> list[Chunk]:
    ext      = path.suffix.lower()
    modality = SUPPORTED.get(ext)
    if modality is None:
        raise ValueError(f"Unsupported extension: {ext}")

    fid = file_id or str(uuid.uuid4())[:8]
    log.info("Ingesting %s [%s]", path.name, ext)

    if ext == ".pdf":
        raw = await _parse_pdf(path)
    elif ext in (".docx", ".doc"):
        raw = await _parse_docx(path)
    elif ext == ".txt":
        raw = await _parse_txt(path)
    elif modality == Modality.image:
        raw = await _parse_image(path)
    else:
        raw = await _parse_audio(path)

    chunks = _to_chunks(raw, path, modality, fid)
    log.info("  -> %d chunks from %s", len(chunks), path.name)
    return chunks


# ── PDF ───────────────────────────────────────────────────────
async def _parse_pdf(path: Path) -> list[dict]:
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _pdf_sync, path)

def _pdf_sync(path: Path) -> list[dict]:
    import fitz
    import re
    doc    = fitz.open(str(path))
    pages  = []
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+',  ' ',    text)
        if text:
            pages.append({ "text": text, "page": i + 1 })
    doc.close()
    return pages


# ── DOCX ──────────────────────────────────────────────────────
async def _parse_docx(path: Path) -> list[dict]:
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _docx_sync, path)

def _docx_sync(path: Path) -> list[dict]:
    from docx import Document
    doc      = Document(str(path))
    sections = []
    cur_head = "Introduction"
    buf      = []
    section_num = 1   # ← use section number as "page"

    for para in doc.paragraphs:
        style = para.style.name.lower()
        text  = para.text.strip()
        if not text:
            continue
        if style.startswith("heading"):
            if buf:
                sections.append({
                    "text":    " ".join(buf),
                    "heading": cur_head,
                    "page":    section_num,   # ← section as page
                })
                buf = []
                section_num += 1
            cur_head = text
        else:
            buf.append(text)

    if buf:
        sections.append({
            "text":    " ".join(buf),
            "heading": cur_head,
            "page":    section_num,
        })

    for table in doc.tables:
        rows     = [
            " | ".join(c.text.strip() for c in row.cells)
            for row in table.rows
        ]
        tbl_text = "\n".join(r for r in rows if r.strip())
        if tbl_text:
            section_num += 1
            sections.append({
                "text":    tbl_text,
                "heading": "Table",
                "page":    section_num,
            })

    return sections

# ── TXT ───────────────────────────────────────────────────────
async def _parse_txt(path: Path) -> list[dict]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return [{ "text": text, "page": None }]


# ── IMAGE ─────────────────────────────────────────────────────
async def _parse_image(path: Path) -> list[dict]:
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _image_sync, path)

def _image_sync(path: Path) -> list[dict]:
    import easyocr
    from PIL import Image
    from config import settings

    img = Image.open(path).convert("RGB")
    w, h = img.size

    # resize large images for speed
    if max(w, h) > 1280:
        img.thumbnail((1280, 1280), Image.LANCZOS)
        w, h = img.size

    reader   = easyocr.Reader(
        ["en"], gpu=False,
        model_storage_directory=str(settings.weights_dir),
        download_enabled=False,
    )
    results  = reader.readtext(str(path), detail=0, paragraph=True)
    ocr_text = " ".join(results).strip()
    desc     = f"[IMAGE {w}×{h}px] {path.stem}."
    if ocr_text:
        desc += f" Text: {ocr_text[:300]}"
    else:
        desc += " No text detected."

    return [{ "text": desc, "page": None, "ocr": ocr_text }]


# ── AUDIO ─────────────────────────────────────────────────────
async def _parse_audio(path: Path) -> list[dict]:
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _audio_sync, path)

def _audio_sync(path: Path) -> list[dict]:
    import whisper, subprocess, tempfile
    from config import settings

    # convert to 16kHz mono WAV
    if path.suffix.lower() != ".wav":
        wav = Path(tempfile.mktemp(suffix=".wav"))
        subprocess.run(
            ["ffmpeg", "-y", "-i", str(path),
             "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le", str(wav)],
            check=True, capture_output=True,
        )
    else:
        wav = path

    model  = whisper.load_model(
        "tiny.en",
        download_root=str(settings.weights_dir),
    )
    result = model.transcribe(
        str(wav), verbose=False, language="en",
        fp16=False, condition_on_previous_text=True,
    )

    if wav != path:
        wav.unlink(missing_ok=True)

    segs = []
    for s in result.get("segments", []):
        text = s["text"].strip()
        if text and len(text) >= 10:
            segs.append({
                "text":      text,
                "timestamp": round(s["start"], 2),
                "duration":  round(result.get("duration", 0), 1),
                "page":      None,
            })
    return segs


# ── Chunker ───────────────────────────────────────────────────
def _to_chunks(
    raw:      list[dict],
    path:     Path,
    modality: Modality,
    file_id:  str,
) -> list[Chunk]:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    import tiktoken

    enc = tiktoken.get_encoding("cl100k_base")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size        = 512,
        chunk_overlap     = 64,
        length_function   = lambda t: len(enc.encode(t)),
        separators        = ["\n\n", "\n", ". ", "! ", "? ", " ", ""],
        keep_separator    = True,
    )

    chunks = []
    for item in raw:
        text = item.get("text", "").strip()
        if not text:
            continue

        pieces = splitter.split_text(text)
        for piece in pieces:
            if not piece.strip():
                continue
            chunks.append(Chunk(
                chunk_id    = str(uuid.uuid4()),
                text        = piece,
                source_path = str(path),
                modality    = modality,
                page        = item.get("page"),
                timestamp   = item.get("timestamp"),
                heading     = item.get("heading"),
                token_count = len(enc.encode(piece)),
                file_id     = file_id,
            ))
    return chunks
