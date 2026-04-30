from docx import Document
from docx.oxml.ns import qn
import asyncio
from pathlib import Path
from dataclasses import dataclass

@dataclass
class RawSection:
    heading:  str
    text:     str
    source:   str

async def parse_docx(path: Path):
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _extract_sync, path)

def _extract_sync(path: Path):
    doc = Document(str(path))
    sections   = []
    cur_head   = "Introduction"
    buf        = []

    for para in doc.paragraphs:
        style = para.style.name.lower()
        text  = para.text.strip()
        if not text:
            continue

        if style.startswith("heading"):
            # Save current buffer as section before new heading
            if buf:
                sections.append(RawSection(
                    heading=cur_head,
                    text=" ".join(buf),
                    source=str(path)
                ))
                buf = []
            cur_head = text
        else:
            buf.append(text)

    if buf:   # flush remaining
        sections.append(RawSection(heading=cur_head, text=" ".join(buf), source=str(path)))

    # ── Tables ──────────────────────────────────────────────
    for table in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        table_text = "\n".join(" | ".join(r) for r in rows if any(r))
        if table_text:
            sections.append(RawSection(heading="Table", text=table_text, source=str(path)))

    return sections
